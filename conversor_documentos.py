#!/usr/bin/env python3
"""
Conversor de Documentos - Aplicación Web Local
Convierte entre: TXT, DOCX, PDF, RTF, HTML, ODT
"""

import os
import sys
import json
import base64
import tempfile
import threading
import webbrowser
import mimetypes
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import parse_qs, urlparse
from io import BytesIO

# ── Librerías de conversión ──────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import html
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
import pdfplumber
from odf.opendocument import OpenDocumentText, load as odt_load
from odf.style import Style, TextProperties, ParagraphProperties
from odf.text import P, Span
from odf import teletype
import pypdf

# ────────────────────────────────────────────────────────────────────────────
#  CONVERSORES
# ────────────────────────────────────────────────────────────────────────────

def extract_text_from_file(data: bytes, ext: str) -> str:
    """Extrae texto plano de cualquier formato."""
    ext = ext.lower().lstrip(".")

    if ext == "txt":
        return data.decode("utf-8", errors="replace")

    if ext in ("doc", "docx"):
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)

    if ext == "pdf":
        with pdfplumber.open(BytesIO(data)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    if ext == "rtf":
        text = data.decode("utf-8", errors="replace")
        # Quitar bloques de control RTF básicos
        import re
        text = re.sub(r'\{\\[^}]*\}', '', text)
        text = re.sub(r'\\[a-z]+\d* ?', '', text)
        text = re.sub(r'[{}]', '', text)
        return text.strip()

    if ext in ("html", "htm"):
        soup = BeautifulSoup(data.decode("utf-8", errors="replace"), "lxml")
        return soup.get_text(separator="\n")

    if ext == "odt":
        doc = odt_load(BytesIO(data))
        return teletype.extractText(doc.text)

    raise ValueError(f"Formato de entrada no soportado: {ext}")


def extract_html_from_file(data: bytes, ext: str) -> str:
    """Extrae HTML con formato básico preservado."""
    ext = ext.lower().lstrip(".")

    if ext in ("html", "htm"):
        return data.decode("utf-8", errors="replace")

    if ext == "txt":
        text = data.decode("utf-8", errors="replace")
        lines = [html.escape(l) for l in text.splitlines()]
        body = "<br>\n".join(lines)
        return _wrap_html(body)

    if ext in ("doc", "docx"):
        doc = Document(BytesIO(data))
        parts = []
        for p in doc.paragraphs:
            style = p.style.name.lower()
            text_esc = html.escape(p.text)
            if "heading 1" in style:
                parts.append(f"<h1>{text_esc}</h1>")
            elif "heading 2" in style:
                parts.append(f"<h2>{text_esc}</h2>")
            elif "heading 3" in style:
                parts.append(f"<h3>{text_esc}</h3>")
            else:
                runs_html = ""
                for run in p.runs:
                    r = html.escape(run.text)
                    if run.bold:   r = f"<strong>{r}</strong>"
                    if run.italic: r = f"<em>{r}</em>"
                    runs_html += r
                parts.append(f"<p>{runs_html}</p>")
        return _wrap_html("\n".join(parts))

    if ext == "pdf":
        text = extract_text_from_file(data, ext)
        lines = [f"<p>{html.escape(l)}</p>" for l in text.splitlines() if l.strip()]
        return _wrap_html("\n".join(lines))

    if ext == "rtf":
        text = extract_text_from_file(data, ext)
        lines = [f"<p>{html.escape(l)}</p>" for l in text.splitlines() if l.strip()]
        return _wrap_html("\n".join(lines))

    if ext == "odt":
        text = extract_text_from_file(data, ext)
        lines = [f"<p>{html.escape(l)}</p>" for l in text.splitlines() if l.strip()]
        return _wrap_html("\n".join(lines))

    raise ValueError(f"Formato de entrada no soportado: {ext}")


def _wrap_html(body: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="es">
<head><meta charset="UTF-8"><title>Documento Convertido</title>
<style>
  body {{ font-family: Georgia, serif; max-width: 800px; margin: 40px auto;
         line-height: 1.7; color: #222; padding: 0 20px; }}
  h1,h2,h3 {{ color: #1a1a2e; }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def convert_to_txt(data: bytes, src_ext: str) -> bytes:
    return extract_text_from_file(data, src_ext).encode("utf-8")


def convert_to_html(data: bytes, src_ext: str) -> bytes:
    return extract_html_from_file(data, src_ext).encode("utf-8")


def convert_to_docx(data: bytes, src_ext: str) -> bytes:
    text = extract_text_from_file(data, src_ext)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for line in text.splitlines():
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert_to_pdf(data: bytes, src_ext: str) -> bytes:
    text = extract_text_from_file(data, src_ext)
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=11,
        leading=16, alignment=TA_LEFT, spaceAfter=6
    )
    story = []
    for line in text.splitlines():
        if line.strip():
            story.append(Paragraph(html.escape(line), body_style))
        else:
            story.append(Spacer(1, 8))
    if not story:
        story.append(Paragraph("(Documento vacío)", body_style))
    doc.build(story)
    return buf.getvalue()


def convert_to_rtf(data: bytes, src_ext: str) -> bytes:
    text = extract_text_from_file(data, src_ext)
    lines = []
    lines.append(r"{\rtf1\ansi\deff0")
    lines.append(r"{\fonttbl{\f0 Times New Roman;}}")
    lines.append(r"\f0\fs24")
    for line in text.splitlines():
        escaped = line.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
        lines.append(escaped + r"\par")
    lines.append("}")
    return "\n".join(lines).encode("latin-1", errors="replace")


def convert_to_odt(data: bytes, src_ext: str) -> bytes:
    text = extract_text_from_file(data, src_ext)
    doc = OpenDocumentText()
    s = Style(name="Default", family="paragraph")
    s.addElement(TextProperties(fontsize="11pt", fontfamily="Liberation Serif"))
    doc.styles.addElement(s)
    for line in text.splitlines():
        p = P(stylename="Default", text=line)
        doc.text.addElement(p)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Tabla de conversores ─────────────────────────────────────────────────────
CONVERTERS = {
    "txt":  convert_to_txt,
    "docx": convert_to_docx,
    "pdf":  convert_to_pdf,
    "rtf":  convert_to_rtf,
    "html": convert_to_html,
    "odt":  convert_to_odt,
}

MIME_TYPES = {
    "txt":  "text/plain",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf":  "application/pdf",
    "rtf":  "application/rtf",
    "html": "text/html",
    "odt":  "application/vnd.oasis.opendocument.text",
}

# ────────────────────────────────────────────────────────────────────────────
#  HTML DE LA INTERFAZ
# ────────────────────────────────────────────────────────────────────────────
HTML_UI = r"""<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Conversor de Documentos</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
  <link href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.min.css" rel="stylesheet">
  <style>
    :root {
      --bg: #f0f2f7;
      --card: #ffffff;
      --accent: #4361ee;
      --accent2: #7209b7;
      --text: #1c1c2e;
      --muted: #6b7280;
      --border: #e2e5f0;
      --success: #06d6a0;
      --radius: 16px;
    }
    * { box-sizing: border-box; }
    body {
      font-family: 'DM Sans', sans-serif;
      background: var(--bg);
      color: var(--text);
      min-height: 100vh;
    }

    /* ── NAV ── */
    .navbar {
      background: var(--card);
      border-bottom: 1px solid var(--border);
      padding: 1rem 2rem;
    }
    .navbar-brand {
      font-weight: 600;
      font-size: 1.1rem;
      color: var(--accent) !important;
      letter-spacing: -0.02em;
    }
    .nav-badge {
      background: linear-gradient(135deg, var(--accent), var(--accent2));
      color: white;
      font-size: .65rem;
      padding: 2px 8px;
      border-radius: 20px;
      font-weight: 500;
      letter-spacing: .04em;
      margin-left: 8px;
      vertical-align: middle;
    }

    /* ── HERO ── */
    .hero {
      background: linear-gradient(135deg, var(--accent) 0%, var(--accent2) 100%);
      padding: 3.5rem 0 5rem;
      position: relative;
      overflow: hidden;
    }
    .hero::after {
      content: '';
      position: absolute;
      bottom: -2px; left: 0; right: 0;
      height: 60px;
      background: var(--bg);
      clip-path: ellipse(55% 100% at 50% 100%);
    }
    .hero h1 {
      font-size: 2.4rem;
      font-weight: 600;
      color: white;
      letter-spacing: -0.03em;
      line-height: 1.2;
    }
    .hero p { color: rgba(255,255,255,.75); font-size: 1.05rem; }

    /* ── FORMAT PILLS ── */
    .fmt-grid {
      display: flex; gap: .5rem; flex-wrap: wrap; justify-content: center;
      margin-top: 1.5rem;
    }
    .fmt-pill {
      background: rgba(255,255,255,.15);
      border: 1px solid rgba(255,255,255,.25);
      color: white;
      border-radius: 40px;
      padding: .3rem .9rem;
      font-size: .8rem;
      font-family: 'DM Mono', monospace;
      font-weight: 500;
      backdrop-filter: blur(4px);
    }

    /* ── CARD ── */
    .main-card {
      background: var(--card);
      border-radius: var(--radius);
      border: 1px solid var(--border);
      box-shadow: 0 4px 24px rgba(0,0,0,.06);
      padding: 2rem;
      margin-top: -2.5rem;
      position: relative;
      z-index: 2;
    }

    /* ── DROP ZONE ── */
    #dropzone {
      border: 2px dashed var(--border);
      border-radius: var(--radius);
      padding: 3rem 1.5rem;
      text-align: center;
      cursor: pointer;
      transition: all .25s ease;
      background: #fafbff;
      position: relative;
    }
    #dropzone:hover, #dropzone.drag-over {
      border-color: var(--accent);
      background: #f0f3ff;
    }
    #dropzone .icon-wrap {
      width: 64px; height: 64px;
      background: linear-gradient(135deg, #eef0ff, #f3e8ff);
      border-radius: 16px;
      display: inline-flex; align-items: center; justify-content: center;
      margin-bottom: 1rem;
      font-size: 1.8rem;
      color: var(--accent);
    }
    #dropzone h5 { font-weight: 600; font-size: 1rem; margin-bottom: .25rem; }
    #dropzone p  { color: var(--muted); font-size: .85rem; margin: 0; }
    #fileInput { display: none; }

    /* ── FILE INFO ── */
    #fileInfo {
      display: none;
      background: #f8faff;
      border: 1px solid var(--border);
      border-radius: 12px;
      padding: 1rem 1.25rem;
    }
    .file-icon {
      width: 44px; height: 44px;
      border-radius: 10px;
      display: flex; align-items: center; justify-content: center;
      font-size: 1.2rem; font-weight: 600;
      color: white; flex-shrink: 0;
    }
    .fmt-txt  { background: #4361ee; }
    .fmt-docx { background: #2b6cb0; }
    .fmt-pdf  { background: #e53e3e; }
    .fmt-rtf  { background: #d69e2e; }
    .fmt-html { background: #38a169; }
    .fmt-odt  { background: #805ad5; }

    /* ── FORMAT SELECT ── */
    .fmt-select-grid {
      display: grid;
      grid-template-columns: repeat(3, 1fr);
      gap: .65rem;
    }
    @media(max-width:576px) { .fmt-select-grid { grid-template-columns: repeat(2,1fr); } }

    .fmt-btn {
      border: 2px solid var(--border);
      border-radius: 12px;
      padding: .75rem .5rem;
      cursor: pointer;
      text-align: center;
      transition: all .2s ease;
      background: var(--card);
      position: relative;
      user-select: none;
    }
    .fmt-btn:hover { border-color: var(--accent); background: #f5f7ff; }
    .fmt-btn.active {
      border-color: var(--accent);
      background: linear-gradient(135deg, #eef0ff 0%, #f3e8ff 100%);
    }
    .fmt-btn .ext {
      font-family: 'DM Mono', monospace;
      font-weight: 600;
      font-size: .9rem;
      color: var(--accent);
    }
    .fmt-btn .desc { font-size: .72rem; color: var(--muted); margin-top: .15rem; }
    .fmt-btn .chk {
      position: absolute; top: 6px; right: 8px;
      color: var(--accent); font-size: .85rem;
      display: none;
    }
    .fmt-btn.active .chk { display: block; }
    .fmt-btn.disabled-fmt {
      opacity: .4; cursor: not-allowed; pointer-events: none;
    }

    /* ── BTN CONVERT ── */
    .btn-convert {
      background: linear-gradient(135deg, var(--accent), var(--accent2));
      color: white; border: none;
      border-radius: 12px;
      padding: .85rem 2.5rem;
      font-weight: 600; font-size: 1rem;
      transition: all .25s;
      width: 100%;
    }
    .btn-convert:hover:not(:disabled) {
      transform: translateY(-2px);
      box-shadow: 0 8px 24px rgba(67,97,238,.35);
      color: white;
    }
    .btn-convert:disabled { opacity: .55; cursor: not-allowed; }

    /* ── RESULT ── */
    #resultArea { display: none; }
    .result-card {
      background: linear-gradient(135deg, #f0fff8, #e6fffa);
      border: 1.5px solid var(--success);
      border-radius: var(--radius);
      padding: 1.5rem;
    }
    .btn-download {
      background: var(--success);
      color: white; border: none;
      border-radius: 10px;
      padding: .7rem 1.5rem;
      font-weight: 600;
      transition: all .2s;
    }
    .btn-download:hover { background: #05b889; color: white; transform: translateY(-1px); }

    /* ── STATUS ── */
    #statusMsg {
      display: none;
      border-radius: 10px;
      padding: .75rem 1.1rem;
      font-size: .88rem;
    }

    /* ── SECTION LABELS ── */
    .section-label {
      font-size: .72rem;
      font-weight: 600;
      letter-spacing: .08em;
      text-transform: uppercase;
      color: var(--muted);
      margin-bottom: .65rem;
    }

    /* ── FOOTER ── */
    footer { text-align: center; color: var(--muted); font-size: .8rem; padding: 2rem 0 1.5rem; }
    footer a { color: var(--accent); text-decoration: none; }

    /* ── SPINNER ── */
    .spinner-grow-sm { width: .7rem; height: .7rem; }

    /* ── ARROW ── */
    .arrow-icon {
      font-size: 1.5rem; color: var(--muted);
      display: flex; align-items: center; justify-content: center;
    }
  </style>
</head>
<body>

<!-- NAVBAR -->
<nav class="navbar">
  <div class="container-fluid">
    <a class="navbar-brand" href="#">
      <i class="bi bi-file-earmark-arrow-right-fill me-2"></i>Conversor de Documentos
    </a>
    <span class="nav-badge">LOCAL · PRIVADO</span>
  </div>
</nav>

<!-- HERO -->
<div class="hero text-center">
  <div class="container">
    <h1>Convierte tus documentos<br>sin salir de tu computadora</h1>
    <p class="mt-2">Procesamiento 100% local · Sin subir archivos a la nube · Sin límites de tamaño</p>
    <div class="fmt-grid">
      <span class="fmt-pill">TXT</span>
      <span class="fmt-pill">DOCX</span>
      <span class="fmt-pill">PDF</span>
      <span class="fmt-pill">RTF</span>
      <span class="fmt-pill">HTML</span>
      <span class="fmt-pill">ODT</span>
    </div>
  </div>
</div>

<!-- MAIN -->
<div class="container" style="max-width:700px; padding-bottom: 3rem;">

  <div class="main-card">

    <!-- PASO 1: subir archivo -->
    <p class="section-label"><i class="bi bi-01-circle me-1"></i>Selecciona el archivo de origen</p>
    <div id="dropzone" onclick="document.getElementById('fileInput').click()">
      <div class="icon-wrap"><i class="bi bi-cloud-upload"></i></div>
      <h5>Arrastra tu archivo aquí</h5>
      <p>o haz clic para buscar · TXT, DOCX, PDF, RTF, HTML, ODT</p>
    </div>
    <input type="file" id="fileInput" accept=".txt,.doc,.docx,.pdf,.rtf,.html,.htm,.odt">

    <!-- Info del archivo -->
    <div id="fileInfo" class="mt-3">
      <div class="d-flex align-items-center gap-3">
        <div id="fileIconBadge" class="file-icon">?</div>
        <div class="flex-grow-1 overflow-hidden">
          <div class="fw-600" id="fileName" style="font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis"></div>
          <div class="text-muted" style="font-size:.8rem" id="fileSize"></div>
        </div>
        <button class="btn btn-sm btn-outline-secondary" onclick="resetFile()" title="Quitar archivo">
          <i class="bi bi-x-lg"></i>
        </button>
      </div>
    </div>

    <hr class="my-4" style="border-color:var(--border)">

    <!-- PASO 2: formato destino -->
    <p class="section-label"><i class="bi bi-02-circle me-1"></i>Elige el formato de salida</p>

    <div class="fmt-select-grid" id="fmtGrid">
      <div class="fmt-btn" data-fmt="txt"  onclick="selectFmt(this)">
        <span class="chk"><i class="bi bi-check-circle-fill"></i></span>
        <div class="ext">.TXT</div>
        <div class="desc">Texto plano</div>
      </div>
      <div class="fmt-btn" data-fmt="docx" onclick="selectFmt(this)">
        <span class="chk"><i class="bi bi-check-circle-fill"></i></span>
        <div class="ext">.DOCX</div>
        <div class="desc">Word</div>
      </div>
      <div class="fmt-btn" data-fmt="pdf"  onclick="selectFmt(this)">
        <span class="chk"><i class="bi bi-check-circle-fill"></i></span>
        <div class="ext">.PDF</div>
        <div class="desc">Portátil</div>
      </div>
      <div class="fmt-btn" data-fmt="rtf"  onclick="selectFmt(this)">
        <span class="chk"><i class="bi bi-check-circle-fill"></i></span>
        <div class="ext">.RTF</div>
        <div class="desc">Rich Text</div>
      </div>
      <div class="fmt-btn" data-fmt="html" onclick="selectFmt(this)">
        <span class="chk"><i class="bi bi-check-circle-fill"></i></span>
        <div class="ext">.HTML</div>
        <div class="desc">Página web</div>
      </div>
      <div class="fmt-btn" data-fmt="odt"  onclick="selectFmt(this)">
        <span class="chk"><i class="bi bi-check-circle-fill"></i></span>
        <div class="ext">.ODT</div>
        <div class="desc">LibreOffice</div>
      </div>
    </div>

    <hr class="my-4" style="border-color:var(--border)">

    <!-- STATUS -->
    <div id="statusMsg" class="alert mb-3"></div>

    <!-- BTN -->
    <button class="btn btn-convert" id="btnConvert" disabled onclick="doConvert()">
      <i class="bi bi-arrow-left-right me-2"></i>Convertir documento
    </button>

    <!-- RESULTADO -->
    <div id="resultArea" class="mt-4">
      <div class="result-card d-flex align-items-center gap-3 flex-wrap">
        <div>
          <i class="bi bi-check-circle-fill text-success fs-4"></i>
        </div>
        <div class="flex-grow-1">
          <div class="fw-bold">¡Conversión completada!</div>
          <div class="text-muted" style="font-size:.85rem" id="resultInfo"></div>
        </div>
        <button class="btn btn-download" id="btnDownload">
          <i class="bi bi-download me-2"></i>Descargar
        </button>
      </div>
    </div>

  </div><!-- /main-card -->
</div><!-- /container -->

<footer>
  <div>Conversor de Documentos · Procesamiento 100% local en tu máquina</div>
  <div class="mt-1" style="font-size:.75rem">TXT · DOCX · PDF · RTF · HTML · ODT</div>
</footer>

<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
<script>
  let selectedFile = null;
  let selectedFmt  = null;
  let convertedBlob = null;
  let convertedName = "";

  const fmtColors = {
    txt:  'fmt-txt',  docx: 'fmt-docx',
    pdf:  'fmt-pdf',  rtf:  'fmt-rtf',
    html: 'fmt-html', odt:  'fmt-odt',
    htm:  'fmt-html', doc:  'fmt-docx',
  };

  // ── Drag & drop ─────────────────────────────────────────────────
  const dz = document.getElementById('dropzone');
  dz.addEventListener('dragover', e => { e.preventDefault(); dz.classList.add('drag-over'); });
  dz.addEventListener('dragleave', ()=> dz.classList.remove('drag-over'));
  dz.addEventListener('drop', e => {
    e.preventDefault(); dz.classList.remove('drag-over');
    if (e.dataTransfer.files[0]) loadFile(e.dataTransfer.files[0]);
  });
  document.getElementById('fileInput').addEventListener('change', e => {
    if (e.target.files[0]) loadFile(e.target.files[0]);
  });

  function loadFile(f) {
    const ext = f.name.split('.').pop().toLowerCase();
    const allowed = ['txt','doc','docx','pdf','rtf','html','htm','odt'];
    if (!allowed.includes(ext)) {
      showStatus('Formato de archivo no soportado: .' + ext, 'danger'); return;
    }
    selectedFile = f;
    document.getElementById('dropzone').style.display = 'none';
    const fi = document.getElementById('fileInfo');
    fi.style.display = 'block';
    document.getElementById('fileName').textContent = f.name;
    document.getElementById('fileSize').textContent = formatBytes(f.size);

    const badge = document.getElementById('fileIconBadge');
    badge.textContent = ext.toUpperCase().slice(0,3);
    badge.className = 'file-icon ' + (fmtColors[ext] || 'fmt-txt');

    // Deshabilitar formato igual al origen
    document.querySelectorAll('.fmt-btn').forEach(btn => {
      const srcExt = ext === 'doc' ? 'docx' : ext === 'htm' ? 'html' : ext;
      if (btn.dataset.fmt === srcExt) btn.classList.add('disabled-fmt');
      else btn.classList.remove('disabled-fmt');
    });

    hideStatus();
    document.getElementById('resultArea').style.display = 'none';
    checkReady();
  }

  function resetFile() {
    selectedFile = null;
    document.getElementById('dropzone').style.display = 'block';
    document.getElementById('fileInfo').style.display = 'none';
    document.getElementById('fileInput').value = '';
    document.querySelectorAll('.fmt-btn').forEach(b => b.classList.remove('disabled-fmt'));
    document.getElementById('resultArea').style.display = 'none';
    hideStatus();
    checkReady();
  }

  function selectFmt(el) {
    document.querySelectorAll('.fmt-btn').forEach(b => b.classList.remove('active'));
    el.classList.add('active');
    selectedFmt = el.dataset.fmt;
    checkReady();
  }

  function checkReady() {
    const ok = selectedFile && selectedFmt;
    document.getElementById('btnConvert').disabled = !ok;
  }

  // ── Convertir ────────────────────────────────────────────────────
  async function doConvert() {
    const btn = document.getElementById('btnConvert');
    btn.disabled = true;
    btn.innerHTML = '<span class="spinner-grow spinner-grow-sm me-2"></span>Convirtiendo…';
    hideStatus();
    document.getElementById('resultArea').style.display = 'none';

    try {
      const reader = new FileReader();
      reader.readAsDataURL(selectedFile);
      reader.onload = async () => {
        const b64 = reader.result.split(',')[1];
        const srcExt = selectedFile.name.split('.').pop().toLowerCase();

        const res = await fetch('/convert', {
          method: 'POST',
          headers: {'Content-Type': 'application/json'},
          body: JSON.stringify({ filename: selectedFile.name, src_ext: srcExt, dst_fmt: selectedFmt, data: b64 })
        });

        if (!res.ok) {
          const err = await res.json();
          showStatus('<i class="bi bi-exclamation-triangle-fill me-2"></i>' + (err.error || 'Error desconocido'), 'danger');
        } else {
          const blob = await res.blob();
          const baseName = selectedFile.name.replace(/\.[^/.]+$/, '');
          convertedName = baseName + '.' + selectedFmt;
          convertedBlob = blob;

          document.getElementById('resultInfo').textContent =
            selectedFile.name + ' → ' + convertedName + ' (' + formatBytes(blob.size) + ')';
          document.getElementById('btnDownload').onclick = () => {
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url; a.download = convertedName; a.click();
            URL.revokeObjectURL(url);
          };
          document.getElementById('resultArea').style.display = 'block';
        }

        btn.disabled = false;
        btn.innerHTML = '<i class="bi bi-arrow-left-right me-2"></i>Convertir documento';
      };
    } catch(e) {
      showStatus('Error: ' + e.message, 'danger');
      btn.disabled = false;
      btn.innerHTML = '<i class="bi bi-arrow-left-right me-2"></i>Convertir documento';
    }
  }

  function showStatus(msg, type) {
    const el = document.getElementById('statusMsg');
    el.className = 'alert alert-' + type + ' mb-3';
    el.innerHTML = msg;
    el.style.display = 'block';
  }
  function hideStatus() {
    document.getElementById('statusMsg').style.display = 'none';
  }
  function formatBytes(b) {
    if (b < 1024) return b + ' B';
    if (b < 1048576) return (b/1024).toFixed(1) + ' KB';
    return (b/1048576).toFixed(1) + ' MB';
  }
</script>
</body>
</html>
"""

# ────────────────────────────────────────────────────────────────────────────
#  SERVIDOR HTTP
# ────────────────────────────────────────────────────────────────────────────

class Handler(BaseHTTPRequestHandler):

    def log_message(self, fmt, *args):
        pass  # silenciar logs

    def do_GET(self):
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.end_headers()
        self.wfile.write(HTML_UI.encode("utf-8"))

    def do_POST(self):
        if self.path != "/convert":
            self.send_response(404); self.end_headers(); return

        length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(length)
        try:
            payload = json.loads(body)
            src_ext = payload["src_ext"].lower().lstrip(".")
            dst_fmt = payload["dst_fmt"].lower()
            file_data = base64.b64decode(payload["data"])

            if dst_fmt not in CONVERTERS:
                raise ValueError(f"Formato destino no soportado: {dst_fmt}")

            result = CONVERTERS[dst_fmt](file_data, src_ext)

            mime = MIME_TYPES.get(dst_fmt, "application/octet-stream")
            self.send_response(200)
            self.send_header("Content-Type", mime)
            self.send_header("Content-Length", str(len(result)))
            self.end_headers()
            self.wfile.write(result)

        except Exception as exc:
            err = json.dumps({"error": str(exc)}).encode()
            self.send_response(400)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(err)))
            self.end_headers()
            self.wfile.write(err)


def run_server(port: int = 8765):
    # En Render se usa 0.0.0.0; en local se usa 127.0.0.1
    is_render = "RENDER" in os.environ
    host = "0.0.0.0" if is_render else "127.0.0.1"
    server = HTTPServer((host, port), Handler)
    url = f"http://127.0.0.1:{port}"
    print(f"\n{'─'*50}")
    print(f"  📄  Conversor de Documentos")
    print(f"{'─'*50}")
    if not is_render:
        print(f"  🌐  Abriendo en: {url}")
        print(f"  🔒  Procesamiento 100% local")
        print(f"  ⌨️   Presiona Ctrl+C para salir")
        threading.Timer(1.0, lambda: webbrowser.open(url)).start()
    else:
        print(f"  🌐  Servidor en puerto {port}")
    print(f"{'─'*50}\n")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n✅  Servidor detenido. ¡Hasta luego!")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", sys.argv[1] if len(sys.argv) > 1 else 8765))
    run_server(port)